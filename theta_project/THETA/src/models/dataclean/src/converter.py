"""
Module for converting various text file formats to Word documents.
"""
import os
import docx
import csv
import json
import re
import xml.etree.ElementTree as ET
from pathlib import Path
from tqdm import tqdm
import PyPDF2
from pdfminer.high_level import extract_text as pdfminer_extract_text
from pdf2docx import Converter as PDFToDocxConverter
import pandas as pd


class TextConverter:
    """
    Class for converting various text file formats to Word documents.
    Supports txt, pdf, csv, and other text-based formats.
    """
    
    def __init__(self):
        """Initialize the TextConverter class."""
        self.supported_formats = ['.txt', '.pdf', '.csv', '.json', '.xml', '.html', '.PDF', '.docx', '.xlsx', '.xls']
    
    def is_supported(self, file_path):
        """
        Check if the file format is supported.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            bool: True if the format is supported, False otherwise
        """
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.supported_formats
    
    def extract_text_from_pdf(self, file_path):
        """
        Extract text from a PDF file using multiple methods for better results.
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            str: Extracted text content
        """
        text = ""
        
        # Try using PyPDF2 first
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    text += pdf_reader.pages[page_num].extract_text() + "\n"
        except Exception as e:
            print(f"PyPDF2 extraction failed: {e}")
        
        # If PyPDF2 didn't extract much text, try pdfminer
        if len(text.strip()) < 100:
            try:
                text = pdfminer_extract_text(file_path)
            except Exception as e:
                print(f"pdfminer extraction failed: {e}")
        
        # If text extraction failed, try a more basic approach
        if len(text.strip()) < 100:
            try:
                with open(file_path, 'rb') as file:
                    # Try to extract text directly from binary
                    content = file.read()
                    # Look for text patterns in binary content
                    text_chunks = re.findall(b'[\x20-\x7E]{4,}', content)
                    additional_text = b'\n'.join(text_chunks).decode('utf-8', errors='ignore')
                    text += "\n" + additional_text
            except Exception as e:
                print(f"Basic text extraction failed: {e}")
        
        return text
    
    def extract_text_from_docx(self, file_path):
        """
        Extract text from a DOCX file.
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            str: Extracted text content
        """
        try:
            doc = docx.Document(file_path)
            text = '\n\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            return text
        except Exception as e:
            print(f"Error extracting text from DOCX {file_path}: {e}")
            return ""
    
    def extract_text_from_csv(self, file_path):
        """
        Extract text from a CSV file.
        
        Args:
            file_path (str): Path to the CSV file
            
        Returns:
            str: Extracted text content
        """
        try:
            text = []
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f)
                for row in reader:
                    text.append(' '.join(row))
            return '\n'.join(text)
        except Exception as e:
            print(f"Error extracting text from CSV {file_path}: {e}")
            return ""
    
    def extract_text_from_json(self, file_path):
        """
        Extract text from a JSON file.
        
        Args:
            file_path (str): Path to the JSON file
            
        Returns:
            str: Extracted text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                data = json.load(f)
            
            # Convert JSON to string representation
            if isinstance(data, dict):
                text = '\n'.join([f"{k}: {v}" for k, v in data.items()])
            elif isinstance(data, list):
                text = '\n'.join([str(item) for item in data])
            else:
                text = str(data)
            
            return text
        except Exception as e:
            print(f"Error extracting text from JSON {file_path}: {e}")
            return ""
    
    def extract_text_from_xml(self, file_path):
        """
        Extract text from an XML file.
        
        Args:
            file_path (str): Path to the XML file
            
        Returns:
            str: Extracted text content
        """
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Extract all text from XML elements
            def extract_text_from_element(element):
                text = element.text or ''
                for child in element:
                    text += ' ' + extract_text_from_element(child)
                if element.tail:
                    text += ' ' + element.tail
                return text
            
            return extract_text_from_element(root)
        except Exception as e:
            print(f"Error extracting text from XML {file_path}: {e}")
            return ""
    
    def extract_text_from_html(self, file_path):
        """
        Extract text from an HTML file.
        
        Args:
            file_path (str): Path to the HTML file
            
        Returns:
            str: Extracted text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            # Remove HTML tags using regex
            clean_text = re.sub(r'<[^>]+>', ' ', html_content)
            clean_text = re.sub(r'\s+', ' ', clean_text).strip()
            
            return clean_text
        except Exception as e:
            print(f"Error extracting text from HTML {file_path}: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path):
        """
        Extract text from a TXT file.
        
        Args:
            file_path (str): Path to the TXT file
            
        Returns:
            str: Extracted text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"Error extracting text from TXT {file_path}: {e}")
            return ""
    
    def extract_text_from_xlsx(self, file_path):
        """
        Extract text from an Excel file (.xlsx or .xls).
        
        Args:
            file_path (str): Path to the Excel file
            
        Returns:
            str: Extracted text content
        """
        try:
            df = pd.read_excel(file_path)
            # Convert all cells to string and join
            text_parts = []
            for col in df.columns:
                text_parts.append(str(col))
            for _, row in df.iterrows():
                row_text = ' '.join([str(val) for val in row.values if pd.notna(val)])
                text_parts.append(row_text)
            return '\n'.join(text_parts)
        except Exception as e:
            print(f"Error extracting text from Excel {file_path}: {e}")
            return ""
    
    def extract_text(self, file_path):
        """
        Extract text from the file.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            str: Extracted text content
        """
        # Get file extension
        ext = os.path.splitext(file_path)[1].lower()
        
        # Use appropriate extraction method based on file type
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif ext == '.csv':
            return self.extract_text_from_csv(file_path)
        elif ext == '.json':
            return self.extract_text_from_json(file_path)
        elif ext == '.xml':
            return self.extract_text_from_xml(file_path)
        elif ext == '.html' or ext == '.htm':
            return self.extract_text_from_html(file_path)
        elif ext == '.txt':
            return self.extract_text_from_txt(file_path)
        elif ext in ['.xlsx', '.xls']:
            return self.extract_text_from_xlsx(file_path)
        else:
            # Try to read as plain text for unsupported formats
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            except Exception as e:
                print(f"Error extracting text from {file_path}: {e}")
                return ""
    
    def convert_pdf_to_word(self, file_path, output_path):
        """
        Convert a PDF file to Word format using pdf2docx.
        
        Args:
            file_path (str): Path to the PDF file
            output_path (str): Path to save the Word document
            
        Returns:
            str: Path to the created Word document
        """
        try:
            # Try using pdf2docx for direct PDF to DOCX conversion
            # This preserves formatting better than text extraction
            converter = PDFToDocxConverter(file_path)
            converter.convert(output_path, start=0, end=None)
            converter.close()
            return output_path
        except Exception as e:
            print(f"Direct PDF to Word conversion failed: {e}")
            print("Falling back to text extraction method...")
            
            # Fall back to text extraction method
            text = self.extract_text_from_pdf(file_path)
            
            # Create a new Word document
            doc = docx.Document()
            
            # Split text by paragraphs and add each paragraph
            paragraphs = text.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Save the document
            doc.save(output_path)
            
            return output_path
    
    def convert_to_word(self, file_path, output_path=None):
        """
        Convert a file to Word format.
        
        Args:
            file_path (str): Path to the input file
            output_path (str, optional): Path to save the Word document. 
                                         If None, will use the same name with .docx extension.
                                         
        Returns:
            str: Path to the created Word document
        """
        if not self.is_supported(file_path):
            raise ValueError(f"Unsupported file format: {file_path}")
        
        if output_path is None:
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            output_dir = os.path.dirname(file_path)
            output_path = os.path.join(output_dir, f"{base_name}.docx")
        
        # Check if the file is a PDF
        if file_path.lower().endswith('.pdf'):
            return self.convert_pdf_to_word(file_path, output_path)
        
        # For other file types
        # Extract text from the file
        text = self.extract_text(file_path)
        
        # Create a new Word document
        doc = docx.Document()
        
        # Split text by paragraphs and add each paragraph
        paragraphs = text.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Save the document
        doc.save(output_path)
        
        return output_path
    
    def batch_convert(self, input_dir, output_dir=None, recursive=False):
        """
        Convert all supported files in a directory to Word format.
        
        Args:
            input_dir (str): Directory containing files to convert
            output_dir (str, optional): Directory to save the Word documents.
                                        If None, will use the same directory.
            recursive (bool): Whether to search for files recursively
            
        Returns:
            list: Paths to the created Word documents
        """
        if output_dir is None:
            output_dir = input_dir
        
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        converted_files = []
        
        # Get all files in the directory
        if recursive:
            all_files = [os.path.join(dp, f) for dp, dn, filenames in os.walk(input_dir) 
                        for f in filenames]
        else:
            all_files = [os.path.join(input_dir, f) for f in os.listdir(input_dir) 
                        if os.path.isfile(os.path.join(input_dir, f))]
        
        # Filter supported files
        supported_files = [f for f in all_files if self.is_supported(f)]
        
        # Convert each file
        for file_path in tqdm(supported_files, desc="Converting files"):
            try:
                rel_path = os.path.relpath(file_path, input_dir)
                output_path = os.path.join(output_dir, os.path.splitext(rel_path)[0] + '.docx')
                
                # Create subdirectories if needed
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
                
                # Convert the file
                converted_file = self.convert_to_word(file_path, output_path)
                converted_files.append(converted_file)
            except Exception as e:
                print(f"Error converting {file_path}: {e}")
        
        return converted_files
